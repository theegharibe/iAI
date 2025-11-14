"""Query Manager - Ask questions from uploaded documents with RAG."""

import json
import os
from pathlib import Path
from typing import List, Dict, Optional
from .rag_manager import SimpleRAGManager
from .advanced_fine_tune import AdvancedFineTuneManager


class QueryManager:
    """
    Query Manager for asking questions from indexed documents.
    
    Features:
    - Index documents for fast retrieval
    - Ask questions and get answers from uploaded files
    - Support for multiple documents
    - Context-aware retrieval
    """
    
    def __init__(self, data_dir: str = "data/fine_tune"):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        
        self.index_file = self.data_dir / "query_index.json"
        self.docs_file = self.data_dir / "query_documents.json"
        
        # RAG manager for retrieval
        self.rag = SimpleRAGManager()
        
        # Advanced fine-tune for extraction
        self.extractor = AdvancedFineTuneManager()
        
        # In-memory storage
        self.documents = {}  # {doc_id: {name, content, qa_pairs}}
        self.indexed_docs = []
        
        self._load_documents()
    
    def _load_documents(self):
        """Load existing documents and re-index them."""
        if self.docs_file.exists():
            try:
                with open(self.docs_file, 'r', encoding='utf-8') as f:
                    self.documents = json.load(f)
                    
                # Re-index all loaded documents into RAG
                if self.documents:
                    all_qa_pairs = []
                    for doc in self.documents.values():
                        all_qa_pairs.extend(doc.get('qa_pairs', []))
                    if all_qa_pairs:
                        self.rag.index_documents(all_qa_pairs)
            except Exception as e:
                print(f"Error loading documents: {e}")
    
    def _save_documents(self):
        """Save documents."""
        try:
            with open(self.docs_file, 'w', encoding='utf-8') as f:
                json.dump(self.documents, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving documents: {e}")
    
    def add_document(self, file_path: str, doc_name: Optional[str] = None) -> Dict:
        """
        Add and index a document for querying.
        
        Args:
            file_path: Path to document (TXT, PDF, DOCX, JSON)
            doc_name: Optional custom name
            
        Returns:
            {
                'success': bool,
                'doc_id': str,
                'pairs_extracted': int,
                'time_seconds': float,
                'message': str
            }
        """
        try:
            import time
            start_time = time.time()
            
            file_path = Path(file_path)
            if not file_path.exists():
                return {
                    'success': False,
                    'message': f'File not found: {file_path}'
                }
            
            # Generate doc ID
            doc_id = file_path.stem.replace(' ', '_')
            if doc_id in self.documents:
                doc_id = f"{doc_id}_{len(self.documents)}"
            
            # Extract content based on file type
            content = ""
            if file_path.suffix.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            elif file_path.suffix.lower() == '.pdf':
                content = ""
                
                # Method 1: Try pdfminer first (best for image PDFs)
                try:
                    print(f"Extracting PDF with pdfminer...")
                    from pdfminer.high_level import extract_text
                    content = extract_text(str(file_path))
                    if content and len(content.strip()) > 50:
                        print(f"[OK] pdfminer extracted {len(content)} characters")
                except Exception as e:
                    print(f"pdfminer attempt failed: {str(e)[:100]}")
                    content = ""
                
                # Method 2: Try PyPDF2 if pdfminer failed
                if not content or len(content.strip()) < 50:
                    try:
                        print(f"Trying PyPDF2...")
                        from PyPDF2 import PdfReader
                        reader = PdfReader(file_path)
                        content = ""
                        for i, page in enumerate(reader.pages):
                            try:
                                text = page.extract_text()
                                if text:
                                    content += text + "\n"
                            except Exception:
                                continue
                        if content and len(content.strip()) > 50:
                            print(f"[OK] PyPDF2 extracted {len(content)} characters")
                    except Exception as e:
                        print(f"PyPDF2 attempt failed: {str(e)[:100]}")
                
                # Method 3: Try pdfplumber if previous failed
                if not content or len(content.strip()) < 50:
                    try:
                        print(f"Trying pdfplumber...")
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            content = ""
                            for page in pdf.pages:
                                try:
                                    text = page.extract_text()
                                    if text:
                                        content += text + "\n"
                                except Exception:
                                    continue
                            if content and len(content.strip()) > 50:
                                print(f"[OK] pdfplumber extracted {len(content)} characters")
                    except Exception:
                        pass
                
                # Final check
                if not content or len(content.strip()) < 50:
                    return {'success': False, 'message': 'PDF: No readable text extracted (may be image-only PDF). Try OCR tools.'}
            
            elif file_path.suffix.lower() == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = ""
                    for para in doc.paragraphs:
                        if para.text.strip():
                            content += para.text + "\n"
                    # Also extract from tables if any
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    content += cell.text + "\n"
                except Exception as e:
                    return {'success': False, 'message': f'DOCX read error: {str(e)[:100]}'}
            
            elif file_path.suffix.lower() == '.json':
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        if isinstance(data, list):
                            content = json.dumps(data, ensure_ascii=False)
                        else:
                            content = json.dumps(data, ensure_ascii=False)
                except Exception as e:
                    return {'success': False, 'message': f'JSON read error: {e}'}
            
            else:
                return {'success': False, 'message': f'Unsupported file type: {file_path.suffix}'}
            
            if not content.strip():
                return {'success': False, 'message': 'File is empty'}
            
            # Extract Q&A pairs
            qa_pairs = self.extractor._smart_extract_qa(content, file_path.stem)
            
            if not qa_pairs or len(qa_pairs) == 0:
                # Log what happened
                print(f"Warning: Could not extract Q&A pairs from {file_path.name}")
                print(f"  Content length: {len(content)} characters")
                print(f"  First 200 chars: {content[:200]}")
                # Instead of failing, create a fallback Q&A
                if len(content) > 100:
                    qa_pairs = [{
                        "instruction": f"What does {file_path.stem} say?",
                        "input": "",
                        "output": content[:2000],  # Use first 2000 chars
                        "source": "fallback"
                    }, {
                        "instruction": "Summarize this document",
                        "input": "",
                        "output": content[:2000],
                        "source": "fallback"
                    }]
                else:
                    return {'success': False, 'message': 'Could not extract Q&A pairs'}
            
            # Add document_id to each QA pair for tracking
            for qa in qa_pairs:
                qa['doc_id'] = doc_id  # Track which document this came from
            
            # Store document
            doc_name = doc_name or file_path.name
            self.documents[doc_id] = {
                'name': doc_name,
                'path': str(file_path),
                'content': content[:5000],  # Store first 5000 chars
                'qa_pairs': qa_pairs,
                'total_chars': len(content),
                'pair_count': len(qa_pairs)
            }
            
            # Index for RAG - combine ALL documents
            all_qa_pairs = []
            for doc in self.documents.values():
                all_qa_pairs.extend(doc.get('qa_pairs', []))
            self.rag.index_documents(all_qa_pairs)
            
            self._save_documents()
            
            elapsed = time.time() - start_time
            
            return {
                'success': True,
                'doc_id': doc_id,
                'pairs_extracted': len(qa_pairs),
                'time_seconds': elapsed,
                'message': f'[OK] Indexed {len(qa_pairs)} Q&A pairs in {elapsed:.2f}s'
            }
        
        except Exception as e:
            return {'success': False, 'message': f'Error: {str(e)}'}
    
    def query(self, question: str, doc_id: Optional[str] = None, top_k: int = 3) -> Dict:
        """
        Ask a question and get answers from indexed documents.
        
        Args:
            question: Question to ask
            doc_id: Optional specific document to search
            top_k: Number of top results to return
            
        Returns:
            {
                'success': bool,
                'question': str,
                'answers': [
                    {
                        'rank': int,
                        'text': str,
                        'source': str,
                        'similarity': float,
                        'doc': str
                    },
                    ...
                ],
                'time_seconds': float,
                'message': str
            }
        """
        try:
            import time
            start_time = time.time()
            
            if not question.strip():
                return {'success': False, 'message': 'Question cannot be empty'}
            
            if not self.documents and not self.rag.documents:
                return {
                    'success': False,
                    'message': 'No documents indexed. Please upload a file first.'
                }
            
            # Retrieve using RAG - get more results for filtering
            results = self.rag.retrieve(question, top_k=top_k * 2)
            
            if not results:
                return {
                    'success': True,
                    'question': question,
                    'answers': [],
                    'time_seconds': time.time() - start_time,
                    'message': 'No answers found'
                }
            
            # Format and filter answers - only keep relevant ones
            answers = []
            seen_answers = set()  # Avoid duplicates
            
            for i, result in enumerate(results):
                # Filter by relevance: 0.15 minimum (reasonable threshold)
                similarity = result.get('score', result.get('similarity', 0))
                if similarity < 0.15:
                    continue
                
                # Try to find source document
                doc_name = "Unknown"
                # First try to get doc_id from result
                result_doc_id = result.get('doc_id')
                if result_doc_id and result_doc_id in self.documents:
                    doc_name = self.documents[result_doc_id].get('name', 'Unknown')
                else:
                    # Fallback to old method
                    for doc_id_iter, doc_info in self.documents.items():
                        if doc_id_iter in str(result.get('source', '')):
                            doc_name = doc_info['name']
                            break
                
                # Get the answer text (could be 'answer', 'output', or 'full_text')
                answer_text = (result.get('answer') or 
                              result.get('output') or 
                              result.get('full_text', ''))
                
                # Skip very short or duplicate answers
                if len(answer_text.strip()) < 10:
                    continue
                
                answer_key = answer_text[:100]  # Use first 100 chars as duplicate check
                if answer_key in seen_answers:
                    continue
                
                seen_answers.add(answer_key)
                
                answers.append({
                    'rank': len(answers) + 1,
                    'text': answer_text.strip(),
                    'source': result.get('source', 'Unknown'),
                    'similarity': max(0.0, min(1.0, similarity)),  # Clamp to 0-1
                    'doc': doc_name
                })
                
                # Stop after getting top_k relevant answers
                if len(answers) >= top_k:
                    break
            
            # Sort by similarity descending
            answers.sort(key=lambda x: x['similarity'], reverse=True)
            
            # Re-rank after sorting
            for i, answer in enumerate(answers, 1):
                answer['rank'] = i
            
            elapsed = time.time() - start_time
            
            message = f'Found {len(answers)} relevant answers' if answers else 'No highly relevant answers found'
            
            return {
                'success': True,
                'question': question,
                'answers': answers,
                'time_seconds': elapsed,
                'message': message
            }
        
        except Exception as e:
            return {'success': False, 'message': f'Error: {str(e)}'}
    
    def list_documents(self) -> List[Dict]:
        """List all indexed documents."""
        return [
            {
                'id': doc_id,
                'name': info['name'],
                'pairs': info['pair_count'],
                'chars': info['total_chars']
            }
            for doc_id, info in self.documents.items()
        ]
    
    def remove_document(self, doc_id: str) -> bool:
        """Remove a document from index."""
        if doc_id in self.documents:
            del self.documents[doc_id]
            self._save_documents()
            return True
        return False
    
    def clear_all(self) -> bool:
        """Clear all documents and index."""
        self.documents = {}
        self.rag.documents = []
        self._save_documents()
        return True
    
    def get_stats(self) -> Dict:
        """Get statistics."""
        total_pairs = sum(d['pair_count'] for d in self.documents.values())
        total_chars = sum(d['total_chars'] for d in self.documents.values())
        
        return {
            'documents_count': len(self.documents),
            'total_qa_pairs': total_pairs,
            'total_characters': total_chars,
            'avg_chars_per_doc': total_chars // len(self.documents) if self.documents else 0
        }
